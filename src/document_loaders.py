"""Offline text extraction for PDF, DOCX, TXT, and image files."""

from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

from src.ocr import run_ocr
from src.preprocess import preprocess_input

IMAGE_TYPES = {"image/jpeg", "image/png", "image/webp"}
PDF_TYPE = "application/pdf"
DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TXT_TYPES = {"text/plain", "application/octet-stream"}


@dataclass
class ExtractionResult:
    """Result produced by a local document loader."""

    text: str
    source_type: str
    warnings: list[str] = field(default_factory=list)
    metadata: dict[str, str | int | float | None] = field(default_factory=dict)


def extract_text_from_file(
    file_bytes: bytes,
    filename: str = "",
    content_type: str = "",
) -> ExtractionResult:
    """Extract text locally from supported document formats."""
    suffix = Path(filename).suffix.lower()
    detected_type = _detect_type(suffix, content_type)

    if detected_type == "pdf":
        return _extract_pdf(file_bytes)
    if detected_type == "docx":
        return _extract_docx(file_bytes)
    if detected_type == "txt":
        return _extract_txt(file_bytes)
    if detected_type == "image":
        return _extract_image(file_bytes, content_type)

    raise ValueError("Unsupported file type. Use PDF, DOCX, TXT, JPG, PNG, or WEBP.")


def _detect_type(suffix: str, content_type: str) -> str:
    if content_type == PDF_TYPE or suffix == ".pdf":
        return "pdf"
    if content_type == DOCX_TYPE or suffix == ".docx":
        return "docx"
    if content_type in TXT_TYPES or suffix == ".txt":
        return "txt"
    if content_type in IMAGE_TYPES or suffix in {".jpg", ".jpeg", ".png", ".webp"}:
        return "image"
    return ""


def _extract_pdf(file_bytes: bytes) -> ExtractionResult:
    try:
        import fitz
    except ImportError as exc:
        raise ValueError("PDF extraction requires local dependency PyMuPDF.") from exc

    text_parts: list[str] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        for page in document:
            text_parts.append(page.get_text("text"))

        metadata = {
            "pages": document.page_count,
            "title": document.metadata.get("title") if document.metadata else None,
        }

    return ExtractionResult(
        text="\n".join(text_parts),
        source_type="pdf",
        metadata=metadata,
    )


def _extract_docx(file_bytes: bytes) -> ExtractionResult:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError(
            "DOCX extraction requires local dependency python-docx."
        ) from exc

    document = Document(BytesIO(file_bytes))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_rows: list[str] = []
    for table in document.tables:
        for row in table.rows:
            table_rows.append("\t".join(cell.text for cell in row.cells))

    return ExtractionResult(
        text="\n".join(paragraphs + table_rows),
        source_type="docx",
        metadata={"paragraphs": len(paragraphs), "tables": len(document.tables)},
    )


def _extract_txt(file_bytes: bytes) -> ExtractionResult:
    warnings: list[str] = []
    for encoding in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return ExtractionResult(
                text=file_bytes.decode(encoding),
                source_type="txt",
                warnings=warnings,
                metadata={"encoding": encoding},
            )
        except UnicodeDecodeError:
            warnings.append(f"Could not decode as {encoding}.")

    return ExtractionResult(
        text=file_bytes.decode("utf-8", errors="replace"),
        source_type="txt",
        warnings=["Decoded with replacement characters."],
        metadata={"encoding": "utf-8-replace"},
    )


def _extract_image(file_bytes: bytes, content_type: str) -> ExtractionResult:
    clean_image = preprocess_input(file_bytes, content_type=content_type)
    ocr_result = run_ocr(clean_image)
    warnings = []
    if ocr_result.error:
        warnings.append(ocr_result.error)
    if ocr_result.confidence < 0.60:
        warnings.append("OCR confidence is low; review the extracted JSON.")

    return ExtractionResult(
        text=ocr_result.text,
        source_type="image",
        warnings=warnings,
        metadata={"ocr_confidence": round(ocr_result.confidence, 2)},
    )
